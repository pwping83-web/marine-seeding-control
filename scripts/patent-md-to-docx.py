# -*- coding: utf-8 -*-
"""Markdown(본 레포 특허 제출용)을 단순 변환하여 .docx 생성. 표는 | 로 구분된 행만 처리."""
from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.shared import Pt
except ImportError:
    print("python-docx 필요: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def strip_md(s: str) -> str:
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"`([^`]+)`", r"\1", s)
    s = s.replace(r"\*", "")
    return s.strip()


def is_table_row(line: str) -> bool:
    t = line.strip()
    return t.startswith("|") and t.endswith("|") and t.count("|") >= 2


def parse_table_row(line: str) -> list[str]:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return [strip_md(c) for c in cells]


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(cols):
            cell_text = row[j] if j < len(row) else ""
            table.rows[i].cells[j].text = cell_text
            for p in table.rows[i].cells[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(11)

    lines = text.splitlines()
    i = 0
    in_blockquote = False
    while i < len(lines):
        line = lines[i]
        raw = line.rstrip()

        if raw.strip() == "---":
            i += 1
            continue

        if raw.startswith("> "):
            p = doc.add_paragraph(strip_md(raw[2:]))
            p.paragraph_format.left_indent = Pt(12)
            i += 1
            continue

        if is_table_row(raw):
            table_rows: list[list[str]] = []
            while i < len(lines) and is_table_row(lines[i]):
                row_line = lines[i].strip()
                cells_raw = [c.strip() for c in row_line.strip().strip("|").split("|")]
                if cells_raw and all(re.match(r"^:?-+:?$", c) for c in cells_raw):
                    i += 1
                    continue
                table_rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, table_rows)
            continue

        if raw.startswith("# "):
            doc.add_heading(strip_md(raw[2:]), level=0)
            i += 1
            continue
        if raw.startswith("## "):
            doc.add_heading(strip_md(raw[3:]), level=1)
            i += 1
            continue
        if raw.startswith("### "):
            doc.add_heading(strip_md(raw[4:]), level=2)
            i += 1
            continue

        if not raw.strip():
            i += 1
            continue

        p = doc.add_paragraph(strip_md(raw))
        i += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"Wrote {docx_path}")


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    submit = root / "docs" / "특허" / "변리사-제출용"
    out = submit / "docx"
    pairs = [
        (submit / "01-미팅용-1페이지-발명-제안-요약.md", out / "01-미팅용-1페이지-발명-제안-요약.docx"),
        (submit / "02-발명-제안서-초안-코드베이스정렬.md", out / "02-발명-제안서-초안-코드베이스정렬.docx"),
        (submit / "03-제출물-체크리스트.md", out / "03-제출물-체크리스트.docx"),
    ]
    for md, dx in pairs:
        if not md.exists():
            print(f"Skip (missing): {md}", file=sys.stderr)
            continue
        md_to_docx(md, dx)


if __name__ == "__main__":
    main()
